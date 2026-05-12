"""Generate the Phase 1-A client-info checklist as a Word document (dev items only)."""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

JP_FONT = "Yu Gothic"
ACCENT = RGBColor(0x1F, 0x3A, 0x68)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT_BG = "F2F4F8"
ACCENT_BG = "EEF2F8"


def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)


def apply_jp_font(run, *, size: float | None = None, bold: bool | None = None,
                   color: RGBColor | None = None) -> None:
    run.font.name = JP_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), JP_FONT)
    rFonts.set(qn("w:ascii"), JP_FONT)
    rFonts.set(qn("w:hAnsi"), JP_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text: str, *, size: float = 10.5, bold: bool = False,
                   color: RGBColor | None = None, align=None,
                   space_after: float | None = None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    apply_jp_font(run, size=size, bold=bold, color=color)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)


def add_heading(doc, text: str, level: int) -> None:
    sizes = {1: 16, 2: 12, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    apply_jp_font(run, size=sizes.get(level, 10.5), bold=True, color=ACCENT)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1F3A68")
        pBdr.append(bottom)
        pPr.append(pBdr)


def add_bullets(doc, items: list[str], *, size: float = 10.5) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        apply_jp_font(run, size=size)
        p.paragraph_format.space_after = Pt(2)


def add_table(doc, headers: list[str], rows: list[list[str]], *,
               col_widths: list[float] | None = None,
               first_col_center: bool = True) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    table.autofit = False

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        set_cell_shading(cell, "1F3A68")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        apply_jp_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            if r_idx % 2 == 0:
                set_cell_shading(cell, LIGHT_BG)
            p = cell.paragraphs[0]
            if first_col_center and c_idx == 0 and len(row) > 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            apply_jp_font(run, size=9.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for row in table.rows:
            for c_idx, width in enumerate(col_widths):
                row.cells[c_idx].width = Cm(width)


def add_spacer(doc, pt: float = 4) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)


def add_callout(doc, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Cm(16)
    set_cell_shading(cell, ACCENT_BG)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    apply_jp_font(run, size=10, color=ACCENT, bold=True)
    add_spacer(doc, 4)


def build() -> Path:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = JP_FONT
    style.font.size = Pt(10.5)

    # --- Cover ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    run = title.add_run("Phase 1-A 開発キックオフ")
    apply_jp_font(run, size=24, bold=True, color=ACCENT)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(6)
    run = sub.add_run("ご提供情報チェックリスト")
    apply_jp_font(run, size=16, bold=True, color=GREY)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.paragraph_format.space_before = Pt(4)
    run = sub2.add_run("楽天 × Shopify 在庫・需要予測・自動発注システム")
    apply_jp_font(run, size=11, color=GREY)

    add_spacer(doc, 60)

    info = doc.add_table(rows=4, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("版 数", "v1.0"),
        ("作成日", datetime(2026, 5, 11).strftime("%Y年%m月%d日")),
        ("対 象", "Phase 1-A（取り込み + マスター在庫管理）"),
        ("範 囲", "開発に必要な提供情報のみ"),
    ]
    for r_idx, (k, v) in enumerate(info_data):
        info.rows[r_idx].cells[0].text = ""
        info.rows[r_idx].cells[1].text = ""
        set_cell_shading(info.rows[r_idx].cells[0], "1F3A68")
        p1 = info.rows[r_idx].cells[0].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(k)
        apply_jp_font(r1, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        p2 = info.rows[r_idx].cells[1].paragraphs[0]
        r2 = p2.add_run(v)
        apply_jp_font(r2, size=10)
        info.rows[r_idx].cells[0].width = Cm(4)
        info.rows[r_idx].cells[1].width = Cm(11)

    doc.add_page_break()

    # --- はじめに ---
    add_heading(doc, "はじめに", 1)
    add_paragraph(
        doc,
        "Phase 1-A 開発を円滑に進めるため、開発工程ごとにクライアント様から"
        "ご提供いただきたい情報を整理しました。",
    )
    add_paragraph(
        doc,
        "本書は開発に必要な情報のみを対象とし、契約・法務に関する書面は"
        "スコープ外としています。",
    )
    add_paragraph(
        doc,
        "情報のご提供タイミングは A〜D の4段階に分けています。"
        "特に A. 着手前必須項目 は、W1 開始に直結するため最優先でのご対応を"
        "お願いいたします。",
    )

    add_spacer(doc, 6)
    add_table(
        doc,
        ["区分", "タイミング", "主な内容"],
        [
            ["A", "着手前必須（〜W1開始まで）", "GCP環境 / 楽天RMS API / Shopify Admin API"],
            ["B", "W1（要件確定・設計）", "商品マスタ / 運用情報 / 既存システム連携"],
            ["C", "W2〜W3（実装・テスト）", "テスト環境 / サンプルデータ / 通知・デザイン"],
            ["D", "W4（受入・本番切替）", "検収責任者 / 本番ドメイン / 切替判断"],
        ],
        col_widths=[1.6, 5.0, 9.4],
    )

    add_spacer(doc, 6)
    add_callout(
        doc,
        "★ 楽天RMS API の利用申請が未完了の場合は、申請から1〜2週間を要します。"
        "他項目より先に申請状況のご確認をお願いいたします。",
    )

    doc.add_page_break()

    # --- A. 着手前必須 ---
    add_heading(doc, "A. 着手前必須（〜W1開始まで）", 1)
    add_paragraph(
        doc,
        "GCP環境構築・API疎通確認に必要です。これらが揃わないと W1 を開始できません。",
    )

    add_heading(doc, "A-1. GCP環境", 2)
    add_table(
        doc,
        ["#", "項目", "内容・備考"],
        [
            ["1", "GCPプロジェクトID", "新規作成 または 既存プロジェクトのご指定"],
            ["2", "課金アカウント", "プロジェクトへの紐付けが完了していること"],
            ["3", "利用リージョン", "推奨：asia-northeast1（東京）"],
            ["4", "開発者へのIAM権限付与", "Owner または Editor + Secret Manager Admin"],
            ["5", "既存BigQueryのプロジェクト/データセット名", "日次exportの宛先として"],
            ["6", "組織ポリシー制約の有無", "VPC強制・Public IP禁止等があれば事前共有"],
        ],
        col_widths=[1.2, 6.0, 8.8],
    )

    add_heading(doc, "A-2. 楽天RMS API", 2)
    add_table(
        doc,
        ["#", "項目", "取得方法・備考"],
        [
            ["7", "RMSサービスシークレット（serviceSecret）", "RMS → 拡張サービス → API設定"],
            ["8", "ライセンスキー（licenseKey）", "同上"],
            ["9", "店舗URL / shopUrl", "例：https://www.rakuten.co.jp/yourshop/"],
            ["10", "利用可能なAPI権限", "注文API（getOrder / searchOrder）の権限ON確認"],
            ["11", "楽天ペイ運用切替の有無", "旧楽天注文API / 楽天ペイ注文APIで挙動が異なる"],
            ["12", "API利用申請の状況", "未申請の場合は申請から1〜2週間要するため最優先"],
        ],
        col_widths=[1.2, 6.0, 8.8],
    )

    add_heading(doc, "A-3. Shopify Admin API", 2)
    add_table(
        doc,
        ["#", "項目", "取得方法・備考"],
        [
            ["13", "ストアドメイン", "xxx.myshopify.com 形式"],
            ["14", "Custom App の Admin API access token", "Shopify管理画面 → アプリ → アプリ開発"],
            ["15", "必要スコープの付与", "read_orders / read_products / read_inventory / read_locations / write_inventory（Phase 1-Bで使用）"],
            ["16", "Webhook設定権限", "アプリ作成権限があれば自動設定可"],
            ["17", "API版数の希望", "既定：最新安定版（例 2025-04）"],
        ],
        col_widths=[1.2, 6.0, 8.8],
    )

    doc.add_page_break()

    # --- B. W1 ---
    add_heading(doc, "B. W1（要件確定・設計）で必要", 1)
    add_paragraph(
        doc,
        "設計確定のためのヒアリングが発生します。打ち合わせ1〜2回で収集予定です。",
    )

    add_heading(doc, "B-1. 商品マスタ情報", 2)
    add_table(
        doc,
        ["#", "項目", "形式・用途"],
        [
            ["18", "マスターSKU命名規則", "テキスト or 既存リスト（命名規則確定）"],
            ["19", "既存マスターSKU一覧", "CSV / スプレッドシート（初期投入データ）"],
            ["20", "楽天SKU ⇔ Shopify variant ID 対応表", "CSV（あれば。マッピング初期データ）"],
            ["21", "JANコード一覧", "CSV（あれば。バーコード照合用）"],
            ["22", "商品カテゴリ・属性体系", "テキスト（Phase 2分析準備）"],
            ["23", "想定SKU件数（現状・1年後）", "概算（性能設計の根拠）"],
        ],
        col_widths=[1.2, 6.0, 8.8],
    )
    add_paragraph(
        doc,
        "※ 既存情報がない場合は弊社で雛形提案いたします（運用都合での命名規則変更も可能です）。",
        size=9, color=GREY,
    )

    add_heading(doc, "B-2. 運用情報", 2)
    add_table(
        doc,
        ["#", "項目", "内容"],
        [
            ["24", "管理画面ログインユーザー一覧", "氏名・メールアドレス・役割"],
            ["25", "想定同時利用人数", "性能想定の根拠"],
            ["26", "エラー通知先メールアドレス", "障害通知・日次レポート宛先"],
            ["27", "営業時間・受注ピーク時間帯", "スケール設計・メンテ時間決定"],
            ["28", "現状の在庫同期手順", "クロスモールの何をどう使っているか"],
            ["29", "倉庫・出荷拠点", "単一前提だが念のため確認"],
            ["30", "キャンセル・返品の業務フロー", "補償イベントの仕様確定"],
        ],
        col_widths=[1.2, 6.0, 8.8],
    )

    add_heading(doc, "B-3. 既存システム連携", 2)
    add_table(
        doc,
        ["#", "項目", "内容"],
        [
            ["31", "クロスモール画面のスクリーンショット（5〜10枚）または短時間の画面共有", "現状機能の把握"],
            ["32", "既存BigQueryのスキーマ", "既存テーブルとの整合性確認"],
            ["33", "既存社内システム（ERP・基幹）との連携要否", "Phase 1-Aスコープ外確認"],
        ],
        col_widths=[1.2, 6.0, 8.8],
    )

    doc.add_page_break()

    # --- C. W2-W3 ---
    add_heading(doc, "C. W2〜W3（実装・テスト）で必要", 1)
    add_paragraph(
        doc,
        "実装中に発生する確認事項です。チャットでの随時対応で十分です。",
    )
    add_table(
        doc,
        ["#", "項目", "内容"],
        [
            ["34", "楽天のテスト店舗の有無", "あればテストAPI接続情報、なければ本番店舗で読み取り限定のテスト"],
            ["35", "Shopify Development Store", "テスト用ストア。Partner経由で無償発行可能"],
            ["36", "テスト用注文の作成可否", "楽天本番でテスト注文を1〜2件作成可能か"],
            ["37", "サンプル注文データ", "過去30日分の注文を1ファイル（CSV / JSON）"],
            ["38", "通知文言の確認", "エラー通知・日次レポートの文面・粒度"],
            ["39", "管理画面のロゴ・カラー指定", "任意。指定なければ標準デザイン"],
        ],
        col_widths=[1.2, 6.0, 8.8],
    )

    # --- D. W4 ---
    add_heading(doc, "D. W4（受入・本番切替）で必要", 1)
    add_table(
        doc,
        ["#", "項目", "内容"],
        [
            ["40", "検収責任者", "受入試験の判断者"],
            ["41", "本番ドメイン（管理画面用）", "例 inventory.example.co.jp（DNS設定はクライアント側）"],
            ["42", "SSL証明書方針", "Google managed SSL利用で問題ないか"],
            ["43", "管理画面のIP制限要否", "オフィスIPからのみ等の要件があれば"],
            ["44", "本番Webhook登録のタイミング", "切替日の合意"],
            ["45", "初期在庫データの投入方針", "クロスモール現在値スナップショット または棚卸結果"],
            ["46", "並行稼働期間の運用ルール", "クロスモールと本システムの責務分担"],
        ],
        col_widths=[1.2, 6.0, 8.8],
    )

    doc.add_page_break()

    # --- 認証情報受け渡し ---
    add_heading(doc, "認証情報の受け渡し方法（推奨）", 1)
    add_paragraph(
        doc,
        "セキュリティ上、API認証情報・各種シークレットは以下のいずれかの方法で"
        "お受け渡しをお願いいたします。",
    )
    add_table(
        doc,
        ["推奨度", "方法"],
        [
            ["★★★", "クライアント様ご自身が Secret Manager に直接登録、IAM権限のみ弊社に付与（流出リスク最小）"],
            ["★★☆", "パスワード保護PDF（パスワードは別チャネル送付）"],
            ["★☆☆", "1Password / Bitwarden 等の共有リンク（一時アクセス）"],
            ["非推奨", "平文メール・チャット直貼り"],
        ],
        col_widths=[2.5, 13.5],
        first_col_center=True,
    )

    add_spacer(doc, 4)
    add_callout(
        doc,
        "特に楽天RMSのライセンスキーは流出時の影響が大きいため、"
        "Secret Manager 直接登録方式を強く推奨いたします。"
        "弊社が登録代行する場合は、登録後にクライアント様側でローテーション可能な手順をご案内します。",
    )

    # --- 最優先 ---
    add_heading(doc, "最優先でご対応いただきたい項目", 1)
    add_paragraph(
        doc,
        "下記3点が揃った時点で W1 着手可能となります。",
    )
    add_table(
        doc,
        ["優先度", "項目", "理由"],
        [
            ["1", "楽天RMS API の利用申請状況確認（#12）", "未申請なら1〜2週間要するため最優先"],
            ["2", "GCPプロジェクトの作成と弊社へのIAM権限付与（#1, #4）", "GCPがないと環境構築できません"],
            ["3", "Shopify Custom App の作成と access token 発行（#14）", "概ね30分作業"],
        ],
        col_widths=[1.8, 8.5, 5.7],
    )

    add_spacer(doc, 8)
    add_paragraph(
        doc,
        "ご不明点がございましたらお気軽にお問い合わせください。",
    )
    add_paragraph(
        doc,
        "本書に記載のない事項についても、開発進行上必要となった場合は"
        "その都度ご相談させていただきます。",
        size=9.5, color=GREY,
    )

    add_spacer(doc, 20)
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end.add_run("— 以上 —")
    apply_jp_font(run, size=10.5, color=GREY)

    out_path = Path(__file__).parent / "Phase1-A_ご提供情報チェックリスト.docx"
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    path = build()
    print(f"Generated: {path}")
