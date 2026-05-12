"""Generate a client-friendly Word document for the Phase 1-A requirements definition."""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

JP_FONT = "Yu Gothic"
JP_FONT_HEAVY = "Yu Gothic"
ACCENT = RGBColor(0x1F, 0x3A, 0x68)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT_BG = "F2F4F8"


def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        from docx.oxml import OxmlElement

        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)


def apply_jp_font(run, *, size: float | None = None, bold: bool | None = None,
                   color: RGBColor | None = None) -> None:
    run.font.name = JP_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement

        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), JP_FONT)
    rFonts.set(qn("w:ascii"), JP_FONT)
    rFonts.set(qn("w:hAnsi"), JP_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text: str, *, size: float = 10.5, bold: bool = False,
                   color: RGBColor | None = None, align=None,
                   space_after: float | None = None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    apply_jp_font(run, size=size, bold=bold, color=color)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)


def add_heading(doc, text: str, level: int) -> None:
    sizes = {1: 18, 2: 14, 3: 12, 4: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        p.paragraph_format.page_break_before = False
    run = p.add_run(text)
    apply_jp_font(run, size=sizes.get(level, 10.5), bold=True, color=ACCENT)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        from docx.oxml import OxmlElement

        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1F3A68")
        pBdr.append(bottom)
        pPr.append(pBdr)


def add_bullets(doc, items: list[str], *, size: float = 10.5) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        apply_jp_font(run, size=size)
        p.paragraph_format.space_after = Pt(2)


def add_table(doc, headers: list[str], rows: list[list[str]], *,
               col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    table.autofit = False

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        set_cell_shading(cell, "1F3A68")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        apply_jp_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            if r_idx % 2 == 0:
                set_cell_shading(cell, LIGHT_BG)
            p = cell.paragraphs[0]
            run = p.add_run(value)
            apply_jp_font(run, size=9.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for row in table.rows:
            for c_idx, width in enumerate(col_widths):
                row.cells[c_idx].width = Cm(width)


def add_spacer(doc, pt: float = 4) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)


def add_callout(doc, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "EEF2F8")
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    apply_jp_font(run, size=10, color=ACCENT, bold=True)
    add_spacer(doc, 4)


def build() -> Path:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = JP_FONT
    style.font.size = Pt(10.5)

    # --- 表紙 ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    run = title.add_run("要 件 定 義 書")
    apply_jp_font(run, size=28, bold=True, color=ACCENT)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(8)
    run = sub.add_run("楽天 × Shopify 在庫・需要予測・自動発注システム")
    apply_jp_font(run, size=14, bold=True, color=GREY)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.paragraph_format.space_before = Pt(4)
    run = sub2.add_run("Phase 1-A：取り込み + マスター在庫管理")
    apply_jp_font(run, size=12, color=GREY)

    add_spacer(doc, 60)

    info = doc.add_table(rows=4, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("版 数", "v1.0"),
        ("作成日", datetime(2026, 5, 5).strftime("%Y年%m月%d日")),
        ("対象範囲", "Phase 1-A 詳細仕様 / 全Phase 概要"),
        ("ステータス", "クライアント確認用"),
    ]
    for r_idx, (k, v) in enumerate(info_data):
        info.rows[r_idx].cells[0].text = ""
        info.rows[r_idx].cells[1].text = ""
        set_cell_shading(info.rows[r_idx].cells[0], "1F3A68")
        p1 = info.rows[r_idx].cells[0].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(k)
        apply_jp_font(r1, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        p2 = info.rows[r_idx].cells[1].paragraphs[0]
        r2 = p2.add_run(v)
        apply_jp_font(r2, size=10)
        info.rows[r_idx].cells[0].width = Cm(4)
        info.rows[r_idx].cells[1].width = Cm(10)

    doc.add_page_break()

    # --- 目次（手書き） ---
    add_heading(doc, "目次", 1)
    toc = [
        "1. はじめに",
        "2. プロジェクト概要",
        "3. 段階開発の全体像",
        "4. Phase 1-A スコープ",
        "5. 機能要件",
        "6. 非機能要件",
        "7. 技術スタック",
        "8. データ管理方針",
        "9. 受入基準",
        "10. 開発体制・スケジュール",
        "11. 前提・制約・リスク",
        "12. 変更管理",
        "13. 用語集",
        "14. 次工程",
    ]
    for item in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        apply_jp_font(run, size=11)

    doc.add_page_break()

    # --- 1. はじめに ---
    add_heading(doc, "1. はじめに", 1)
    add_paragraph(
        doc,
        "本書は、楽天市場およびShopifyを中心とした自社EC運営における、"
        "在庫管理・需要予測・自動発注を統合するシステムについて、"
        "Phase 1-A（取り込み + マスター在庫管理）の要件を確定するものです。",
    )
    add_paragraph(
        doc,
        "Phase 1-Aの確実な完成が後続フェーズ（Phase 1-B / Phase 2 / Phase 3、"
        "Amazon連携、卸販売取り込み）の前提となるため、"
        "本書ではスコープ・成果物・受入基準を明確化し、"
        "後続拡張を見据えた設計境界もあわせて定義します。",
    )
    add_callout(
        doc,
        "本書は合意後にベースライン化され、以降の変更は変更管理プロセスにより管理します。",
    )

    # --- 2. プロジェクト概要 ---
    add_heading(doc, "2. プロジェクト概要", 1)
    add_heading(doc, "2.1 背景", 2)
    add_bullets(doc, [
        "現状、楽天とShopify間の在庫同期は手動運用となっており、運用負荷と機会損失リスクが存在する。",
        "既存SaaS（クロスモール等）に依存せず、自社資産としてのデータ基盤を構築したい。",
        "蓄積データを将来のAI需要予測・自動発注の基礎としたい。",
    ])

    add_heading(doc, "2.2 最終ゴール", 2)
    add_bullets(doc, [
        "全チャネル（楽天 / Shopify / Amazon / 卸）の注文・在庫データを中央DBに集約する。",
        "中央DBを「唯一の真実（Single Source of Truth）」として、各チャネルへ在庫を双方向同期する。",
        "蓄積データから需要予測を行い、工場別の発注書ドラフト（PDF + CSV、英/日）を自動生成する。",
        "最終承認・送信は人間が実施し、AIは草案生成までを担当する。",
    ])

    add_heading(doc, "2.3 Phase 1-A の位置づけ", 2)
    add_paragraph(
        doc,
        "Phase 1-A は、最終ゴールに向けた最初の基盤づくりです。"
        "チャネルへの在庫書き戻しは行わず、"
        "「中央DBにデータを正確に集約する」ことに集中します。"
        "これにより、後続フェーズで安全に書き戻し・分析・AI予測へ拡張できる土台を築きます。",
    )

    # --- 3. 段階開発の全体像 ---
    add_heading(doc, "3. 段階開発の全体像", 1)
    add_paragraph(doc, "全フェーズの想定費用・納期は下表のとおりです。")
    add_table(
        doc,
        ["Phase", "内容", "想定費用", "想定納期"],
        [
            ["1-A", "Shopify+楽天 取り込み・中央DB構築（書き戻しなし）", "38万円", "3〜4週間"],
            ["1-B", "Shopify+楽天 書き戻し・双方向同期", "25〜30万円", "3〜4週間"],
            ["2", "在庫分析・欠品予測・発注アラート", "35〜45万円", "4〜5週間"],
            ["3", "AI需要予測・発注書生成", "50〜70万円", "6〜8週間"],
            ["Amazon", "Amazon SP-API連携", "25〜35万円", "3〜4週間"],
            ["卸販売", "CSV / Sheets / 手動入力 取り込み", "15〜25万円", "2〜3週間"],
        ],
        col_widths=[2.0, 7.5, 3.0, 3.0],
    )
    add_paragraph(doc, "全Phase合算概算：188〜243万円", bold=True, color=ACCENT)
    add_paragraph(
        doc,
        "※ Phase 3 は Phase 1〜2 で蓄積されたデータを前提とするため、"
        "最低6ヶ月以上のデータ蓄積後の着手を推奨いたします。",
        size=9.5,
        color=GREY,
    )

    # --- 4. スコープ ---
    add_heading(doc, "4. Phase 1-A スコープ", 1)

    add_heading(doc, "4.1 含まれるもの（IN）", 2)
    add_bullets(doc, [
        "楽天RMS APIによる注文の自動取り込み（Polling方式）",
        "Shopify Admin APIによる注文の自動取り込み（Webhook + Polling冗長化、HMAC検証必須）",
        "中央在庫DB（PostgreSQL on Cloud SQL）の構築 — マスターSKU中心の正規化テーブル設計",
        "SKU単位の在庫減算処理（イベントソーシング型 / 行ロックによる整合性担保）",
        "管理画面 3〜4画面（SKUマッピング / 在庫一覧 / 手動調整 / イベントログ閲覧）",
        "BigQueryへの日次データexport（既存BQ環境を利用）",
        "GCP環境構築（Cloud Run + Cloud Tasks + Secret Manager + Cloud SQL）",
        "基本的な監視・ログ・エラー通知設計",
    ])

    add_heading(doc, "4.2 含まれないもの（OUT）", 2)
    add_paragraph(
        doc,
        "下記は本Phaseの対象外であり、後続フェーズまたは別途見積となります。",
    )
    add_bullets(doc, [
        "中央DBから各チャネルへの在庫書き戻し（→ Phase 1-B）",
        "楽天 ↔ Shopify の双方向自動同期（→ Phase 1-B、当面は既存手段で継続）",
        "同期エラー一覧画面・再実行機能・Slack通知（→ Phase 1-B）",
        "在庫分析ダッシュボード・欠品予測（→ Phase 2）",
        "AI需要予測・発注書生成（→ Phase 3）",
        "Amazon SP-API連携 / 卸販売取り込み（→ 別Phase）",
        "既存システムからの初期データ移行（別途見積）",
        "楽天店舗・Shopifyストアが複数ある場合の追加対応（別途見積：本Phaseは1店舗・1ストア前提）",
        "GCPの月額運用費（クライアント様ご負担、月3,000〜10,000円程度想定）",
    ])

    add_heading(doc, "4.3 前提条件", 2)
    add_bullets(doc, [
        "楽天店舗：1店舗を前提（複数店舗は別途見積）",
        "Shopifyストア：1ストアを前提（複数ストアは別途見積）",
        "倉庫：単一倉庫を前提（複数倉庫対応はPhase 1-B以降で別途）",
        "本Phase期間中、楽天とShopifyの在庫同期は既存の手段で継続いただく（クライアント側責務）",
        "GCPプロジェクト・課金アカウントはクライアント様にご用意いただく",
        "楽天RMS API / Shopify Admin APIの認証情報および店舗権限はクライアント様よりご提供いただく",
        "BigQuery環境は既存のものを使用",
    ])

    # --- 5. 機能要件 ---
    add_heading(doc, "5. 機能要件", 1)

    add_heading(doc, "5.1 楽天注文取り込み（Polling）", 2)
    add_table(
        doc,
        ["No.", "機能", "概要"],
        [
            ["F-RAK-01", "注文一覧定期取得", "楽天RMS APIで一定間隔（既定5分）に新規・更新注文を取得"],
            ["F-RAK-02", "増分取得", "前回成功時刻以降の更新注文のみを取得しAPI消費を最小化"],
            ["F-RAK-03", "注文詳細取得", "注文番号ごとに明細（SKU・数量・金額）を取得"],
            ["F-RAK-04", "重複排除", "(channel, channel_order_id) のUNIQUE制約で重複登録を防止"],
            ["F-RAK-05", "キャンセル処理", "キャンセル状態への遷移を検知し、補償イベントで在庫を戻入"],
            ["F-RAK-06", "リトライ", "一時的失敗（429/5xx）は指数バックオフで最大5回リトライ"],
            ["F-RAK-07", "レート制限遵守", "Token Bucket方式でAPI制限を超えないよう制御"],
            ["F-RAK-08", "認証情報管理", "Secret Managerに格納、コードへの平文埋め込み禁止"],
        ],
        col_widths=[2.5, 4.0, 9.0],
    )

    add_heading(doc, "5.2 Shopify注文取り込み（Webhook + Polling冗長化）", 2)
    add_table(
        doc,
        ["No.", "機能", "概要"],
        [
            ["F-SHO-01", "Webhook受信", "orders/create, orders/updated, orders/cancelled を受信"],
            ["F-SHO-02", "HMAC検証", "X-Shopify-Hmac-Sha256 を必ず検証、失敗時は401で破棄しログ記録"],
            ["F-SHO-03", "即時ACK", "Webhook受信は5秒以内に200応答、処理はCloud Tasksへ非同期投入"],
            ["F-SHO-04", "Polling冗長化", "Webhook欠落に備え、定期Polling（既定15分）で増分取得"],
            ["F-SHO-05", "重複排除", "channel_order_id および Webhook ID で冪等処理"],
            ["F-SHO-06", "キャンセル処理", "楽天と同様、補償イベントで在庫を戻入"],
            ["F-SHO-07", "GraphQL推奨", "Admin APIはGraphQLを優先（コスト効率・将来互換性）"],
            ["F-SHO-08", "認証情報管理", "Secret Manager管理、ローテーション可能な構造"],
        ],
        col_widths=[2.5, 4.0, 9.0],
    )

    add_heading(doc, "5.3 中央在庫DB（マスターSKU中心の正規化）", 2)
    add_paragraph(
        doc,
        "在庫変動の真実を「マスターSKU単位の追記専用イベントログ」として記録します。"
        "現在在庫はイベントの合算から導出され、参照高速化のためのスナップショットを併用します。",
    )
    add_table(
        doc,
        ["No.", "テーブル", "役割"],
        [
            ["F-DB-01", "master_skus", "自社管理の正規SKU。すべての在庫変動の基準"],
            ["F-DB-02", "channel_sku_mappings", "マスターSKU ⇔ 各チャネルSKU/商品IDの対応表"],
            ["F-DB-03", "inventory_events", "在庫変動の追記専用ログ（イベントソーシング）"],
            ["F-DB-04", "inventory_snapshots", "現在在庫キャッシュ。真実はイベント側"],
            ["F-DB-05", "orders / order_items", "各チャネルから取り込んだ注文の正規化表"],
            ["F-DB-06", "mapping_alerts", "未マッピングSKU検出時のアラート管理"],
            ["F-DB-07", "拡張カラム先行用意", "fulfillment_type / marketplace_id を初期から保持"],
            ["F-DB-08", "整合性制約", "UNIQUE/FK/NOT NULL厳格定義、行ロックで並行制御"],
        ],
        col_widths=[2.5, 4.5, 8.5],
    )

    add_heading(doc, "5.4 在庫減算ロジック（イベントソーシング）", 2)
    add_table(
        doc,
        ["No.", "ルール"],
        [
            ["F-INV-01", "注文取り込み時に order_items 単位で order_consumed イベントを追加"],
            ["F-INV-02", "同一注文行からのイベントは UNIQUE 制約で重複登録不可（冪等性）"],
            ["F-INV-03", "未マッピングSKUは減算せず mapping_alerts に登録、注文は保留状態で記録"],
            ["F-INV-04", "キャンセル受信時は cancellation_returned イベントで打ち消し（履歴は削除しない）"],
            ["F-INV-05", "管理画面からの調整は manual_adjust イベント（理由・操作者を記録）"],
            ["F-INV-06", "イベント合算 = スナップショット を日次で自動検証"],
            ["F-INV-07", "注文取込→減算→マッピング解決は単一トランザクション（部分反映禁止）"],
        ],
        col_widths=[2.5, 13.0],
    )

    add_heading(doc, "5.5 管理画面（3〜4画面）", 2)
    add_paragraph(
        doc,
        "実装方式は FastAPI + Jinja2 + Tailwind を第一候補とし、"
        "要件に応じて Retool による代替もご提案いたします。"
        "管理画面は認証必須で、すべての操作はログとして記録されます。",
    )
    add_table(
        doc,
        ["画面", "概要"],
        [
            ["SKUマッピング画面", "マスターSKU ⇔ チャネルSKUの一覧・検索・新規・編集・削除、CSV入出力"],
            ["在庫一覧画面", "マスターSKU単位の現在在庫・直近変動を一覧、SKU検索、低在庫フィルタ"],
            ["手動調整画面", "数量・理由を入力し manual_adjust イベントを発行"],
            ["イベントログ閲覧", "在庫イベントの時系列一覧（SKU・期間・種別・チャネルでフィルタ）"],
        ],
        col_widths=[5.0, 10.5],
    )

    add_heading(doc, "5.6 BigQuery 日次 Export", 2)
    add_bullets(doc, [
        "対象テーブル：orders / order_items / inventory_events / inventory_snapshots / channel_sku_mappings / master_skus",
        "実行：Cloud Scheduler起動で日次1回（例：JST 03:00）",
        "方式：updated_at ベースの増分export、スナップショットは日次フル",
        "パーティション：日付パーティション（event_date / order_date）",
        "リトライ：失敗時は次回起動でキャッチアップ、3日連続失敗で通知",
        "スキーマ管理：テーブル定義をコード管理（Terraform または DDLファイル）",
    ])

    # --- 6. 非機能要件 ---
    add_heading(doc, "6. 非機能要件", 1)

    add_heading(doc, "6.1 性能", 2)
    add_bullets(doc, [
        "Webhook起点の在庫反映：平均5分以内",
        "Polling起点の在庫反映：15分以内",
        "管理画面の通常表示：2秒以内、検索：5秒以内（SKU 5万件想定）",
        "1日あたりの処理想定：注文5,000件 / 在庫イベント20,000件",
    ])

    add_heading(doc, "6.2 可用性・信頼性", 2)
    add_bullets(doc, [
        "Cloud Runのヘルスチェック有効化、自動再起動",
        "Webhook受信は5秒以内に応答、処理はCloud Tasksで非同期化",
        "At-least-once + 冪等性で欠落・重複に耐性",
        "整合性検証バッチを日次実行（イベント合算 = スナップショット）",
    ])

    add_heading(doc, "6.3 セキュリティ", 2)
    add_bullets(doc, [
        "全シークレットはSecret Manager管理、コード/環境変数への直書き禁止",
        "Shopify WebhookのHMAC検証必須",
        "管理画面は認証必須、操作ログを記録",
        "DB通信はSSL強制、Cloud SQLのPrivate IP / IAM認証を検討",
        "最小権限のサービスアカウント設計",
    ])

    add_heading(doc, "6.4 監査・運用", 2)
    add_bullets(doc, [
        "構造化ログ（JSON）をCloud Loggingに出力",
        "エラー時はCloud Logging Alertでメール通知（Slack連携はPhase 1-B）",
        "リクエストID／取込バッチIDによる横断追跡",
        "Cloud SQLの自動バックアップ（日次）",
    ])

    add_heading(doc, "6.5 拡張性（将来Phase対応）", 2)
    add_paragraph(
        doc,
        "将来のチャネル追加（Amazon、卸ほか）に備え、初期段階から下記の拡張ポイントを"
        "コードに組み込みます。これにより、Phase 1-Bや別Phaseの追加開発コストを最小化します。",
    )
    add_bullets(doc, [
        "ChannelAdapter 抽象クラス（fetch_orders / push_inventory / verify_webhook）を共通インターフェースとして導入",
        "Phase 1-Aでは RakutenAdapter / ShopifyAdapter を実装。push_inventory はPhase 1-Bで実装予定",
        "Amazon FBA/MFN対応のため fulfillment_type、複数マーケットプレイス対応のため marketplace_id を初期スキーマに用意",
        "注文ステータスは各チャネル独自定義から「統一ステータス」へAdapter層で正規化",
    ])

    add_heading(doc, "6.6 品質・コーディング規約", 2)
    add_bullets(doc, [
        "Python 3.11+ / 型ヒント必須、ruff + mypyによる静的解析をCI実行",
        "在庫整合性ロジック・冪等性周辺はユニットテスト必須カバレッジ（80%以上）",
        "統合テストは楽天/Shopifyのサンドボックスまたはモックで主要シナリオを実施",
        "AIコーディングツール（Claude Code / Cursor / Copilot）の活用は可。"
        "ただし在庫整合性ロジックは人間レビュー必須、ケース漏れは人間検証、API挙動は実環境で動作確認",
    ])

    # --- 7. 技術スタック ---
    add_heading(doc, "7. 技術スタック", 1)
    add_table(
        doc,
        ["層", "技術", "備考"],
        [
            ["言語", "Python 3.11+", "型ヒント必須"],
            ["Webフレームワーク", "FastAPI", "非同期I/O、OpenAPI自動生成"],
            ["データベース", "PostgreSQL 15+ (Cloud SQL)", "行ロック・UNIQUE制約活用"],
            ["実行基盤", "Cloud Run", "コンテナベース、自動スケール"],
            ["非同期ジョブ", "Cloud Tasks", "Webhook処理・Polling・BQ Export"],
            ["定期実行", "Cloud Scheduler", "Polling・日次バッチ起動"],
            ["シークレット管理", "Secret Manager", "API key / DB認証情報"],
            ["分析基盤", "BigQuery（既存）", "日次exportの宛先"],
            ["管理画面", "FastAPI + Jinja2 + Tailwind", "第一候補。Retoolも代替提案可"],
            ["マイグレーション", "Alembic", "スキーマバージョン管理"],
            ["IaC", "Terraform（推奨）", "GCPリソース定義"],
            ["CI/CD", "GitHub Actions", "テスト・デプロイ自動化"],
        ],
        col_widths=[3.5, 5.5, 6.5],
    )

    # --- 8. データ管理方針 ---
    add_heading(doc, "8. データ管理方針", 1)
    add_paragraph(
        doc,
        "中央DBは「マスターSKU」を中心とした正規化設計を採用します。"
        "これにより、各チャネルでSKU命名が異なっていても統一的に扱え、"
        "1つのマスターSKUに対する複数チャネルSKUの紐付け（セット商品・バリエーション）にも対応します。",
    )
    add_bullets(doc, [
        "在庫変動はすべてマスターSKUを基準に記録",
        "(channel, channel_sku) のUNIQUE制約により重複マッピングを防止",
        "未マッピングSKUは管理画面のアラート画面で検知・解決",
        "Amazonや卸を将来追加する際も、新規マッピング追加のみで既存データを傷めない",
    ])

    add_callout(
        doc,
        "中央DBが「Single Source of Truth（唯一の真実）」となり、"
        "将来のAI需要予測（Phase 3）の基礎となります。",
    )

    # --- 9. 受入基準 ---
    add_heading(doc, "9. 受入基準", 1)
    add_paragraph(
        doc,
        "以下を全て満たすことで Phase 1-A 完了とし、後続Phase着手の判断基準とします。",
    )

    add_heading(doc, "9.1 機能受入", 2)
    add_bullets(doc, [
        "楽天注文1件が平均15分以内に中央DBへ反映され、対応する在庫イベントが記録される",
        "Shopify Webhook経由で注文1件が平均5分以内に在庫反映、HMAC不正は401拒否しログ残存",
        "Webhook欠落時もPollingで24時間以内に必ず取り込まれる",
        "同一注文のWebhookが2回送られても在庫が二重に減らない（冪等性）",
        "キャンセル受信で補償イベントが追加され、在庫スナップショットが回復する",
        "未マッピングSKUの注文はアラート登録、管理画面から解決後に再処理可能",
        "管理画面4画面（マッピング/在庫/手動調整/ログ）が動作する",
        "手動調整は理由と操作者がイベントに記録される",
        "日次でBigQueryに対象テーブルがexportされ、翌朝までに参照可能",
    ])

    add_heading(doc, "9.2 整合性受入", 2)
    add_bullets(doc, [
        "在庫イベント合算とスナップショットが日次自動検証で一致、不一致時は通知",
        "並行注文時も在庫がマイナスにならず、行ロックで直列化される",
        "冪等性UNIQUE制約により、同一イベントの再送で二重登録されない",
    ])

    add_heading(doc, "9.3 非機能受入", 2)
    add_bullets(doc, [
        "全シークレットがSecret Manager管理（リポジトリ内に直書きなし）",
        "GCPリソースがTerraformで再現可能",
        "CIでruff / mypy / pytestが通る",
        "在庫整合性ロジックのユニットテストカバレッジ80%以上",
        "主要ユースケースの統合テスト（モックまたはサンドボックス）が通る",
    ])

    add_heading(doc, "9.4 拡張性受入（将来Phase対応の確認）", 2)
    add_bullets(doc, [
        "ChannelAdapter抽象クラスが実装済み、Phase 1-Bで push_inventory のメソッド実装のみで対応可能",
        "fulfillment_type / marketplace_id カラムがスキーマに存在",
        "新Adapter追加時、コア在庫ロジックの改修が不要であることをコードレビューで確認",
    ])

    add_heading(doc, "9.5 ドキュメント納品物", 2)
    add_bullets(doc, [
        "要件定義書（本書、最終版）",
        "アーキテクチャ設計書",
        "DBスキーマ仕様（DDL + ER図）",
        "API仕様書（OpenAPI 3.0）",
        "管理画面マニュアル",
        "運用手順書（デプロイ・障害時対応・認証情報ローテーション）",
        "ソースコード一式（GitHubリポジトリ）",
    ])

    # --- 10. スケジュール ---
    add_heading(doc, "10. 開発体制・スケジュール", 1)

    add_heading(doc, "10.1 マイルストーン", 2)
    add_table(
        doc,
        ["週", "主な作業", "成果物"],
        [
            ["W1", "要件確定・アーキテクチャ設計・DB設計・GCP環境準備",
             "アーキテクチャ設計書 / DDL / Terraform雛形"],
            ["W2", "Adapter基盤・Shopify取込（Webhook+Polling）・在庫イベント・主要API",
             "コア機能の動作"],
            ["W3", "楽天取込・管理画面・BigQuery export・統合テスト",
             "管理画面4画面 / E2Eテスト通過"],
            ["W4", "サンドボックス→本番疎通・整合性検証・ドキュメント整備・受入試験",
             "全受入基準クリア"],
        ],
        col_widths=[1.5, 7.5, 6.5],
    )

    add_heading(doc, "10.2 コミュニケーション", 2)
    add_bullets(doc, [
        "進捗：週次レポート（GitHub Projects または Notion）",
        "質疑：チャットで随時、判断を要する事項はメールで合意エビデンスを残す",
        "レビュー：節目（W1末・W2末・W3末・W4末）に動作デモを実施",
    ])

    add_heading(doc, "10.3 瑕疵対応", 2)
    add_bullets(doc, [
        "納品後3ヶ月は、受入基準を満たさない不具合を無償対応",
        "仕様変更・運用相談・運用代行は別途月額契約にて対応",
    ])

    # --- 11. リスク ---
    add_heading(doc, "11. 前提・制約・リスク", 1)
    add_table(
        doc,
        ["#", "リスク", "影響", "対応"],
        [
            ["R1", "楽天RMSの特殊在庫（項目選択肢別在庫等）", "取り込み欠落",
             "本Phaseは標準SKU構造を前提、特殊在庫はPhase 1-Bで別途見積"],
            ["R2", "Shopify Webhookの欠落", "在庫ズレ",
             "Polling冗長化により24時間以内に必ず追従"],
            ["R3", "既存運用との並行期間の在庫ズレ", "売り越し",
             "本Phaseは書き戻しなし。中央DBは観測専用、既存同期手段は継続"],
            ["R4", "マッピング作業の運用負荷", "取り込み遅延",
             "アラート画面 + CSV一括登録機能で運用効率化"],
            ["R5", "APIレート制限", "取り込み停止",
             "チャネル別Token Bucket、指数バックオフリトライ"],
            ["R6", "店舗数前提変更（複数店舗化）", "工数増",
             "1店舗・1ストア前提を明記、追加は別途見積"],
            ["R7", "AIコーディング起因のロジック欠陥", "在庫整合性破綻",
             "整合性ロジックは人間レビュー必須、テスト網羅"],
        ],
        col_widths=[1.0, 4.5, 3.0, 7.0],
    )

    # --- 12. 変更管理 ---
    add_heading(doc, "12. 変更管理", 1)
    add_bullets(doc, [
        "本書合意後の要件追加・変更は変更要求書（CR）にて受領",
        "影響範囲・追加工数・納期影響を提示し、書面合意のうえ反映",
        "スコープ外項目（4.2 OUT 列挙）は原則 Phase 1-B 以降に持ち越し",
    ])

    # --- 13. 用語集 ---
    add_heading(doc, "13. 用語集", 1)
    add_table(
        doc,
        ["用語", "意味"],
        [
            ["マスターSKU", "自社管理の正規SKU。全チャネル共通の真実"],
            ["ChannelAdapter", "各チャネル接続を抽象化する共通インターフェース"],
            ["イベントソーシング", "在庫変動を追記専用ログとして記録し、現在状態を導出する設計"],
            ["SSoT", "Single Source of Truth。中央DBが在庫の唯一の真実"],
            ["FBA / MFN", "Amazon の Fulfillment by Amazon / Merchant Fulfilled Network"],
            ["冪等性", "同じ操作を何度実行しても結果が変わらない性質"],
            ["HMAC", "Webhook送信元の真正性を検証する署名アルゴリズム"],
            ["Polling", "定期的にAPIへ問い合わせて差分を取得する方式"],
            ["Webhook", "イベント発生時に外部から通知を受け取る方式"],
        ],
        col_widths=[5.0, 10.5],
    )

    # --- 14. 次工程 ---
    add_heading(doc, "14. 次工程", 1)
    add_paragraph(
        doc,
        "本書のご承認を経て、次は以下のドキュメントを作成のうえ実装に着手いたします。",
    )
    add_bullets(doc, [
        "アーキテクチャ設計書（コンポーネント図 / シーケンス図 / GCPリソース構成）",
        "DBスキーマ仕様書（DDL / ER図 / インデックス / 制約）",
        "API仕様書（OpenAPI 3.0）",
        "画面仕様書（ワイヤーフレーム / 操作フロー）",
    ])
    add_paragraph(
        doc,
        "ご確認のほどよろしくお願い申し上げます。",
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )

    # --- 末尾 ---
    add_spacer(doc, 30)
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end.add_run("— 以上 —")
    apply_jp_font(run, size=10.5, color=GREY)

    out_path = Path(__file__).parent / "01_要件定義書_Phase1-A.docx"
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    path = build()
    print(f"Generated: {path}")
